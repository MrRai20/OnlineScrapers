#!/usr/bin/env python3
r"""
Wiki Encyclopaedia Builder v4 (hardened)
- Structured parsing → Dossier-style outputs
- Formal layout in DOCX / PDF (Unicode font, headers/footers), styled EPUB (optional embedded font)
- EPUBCheck integration (bin\epubcheck.bat preferred; java -cp fallback; else read-back)
- Deterministic ordering, JSON/YAML/CSV, README, search tool, ZIP
"""

from __future__ import annotations
import argparse, csv, html, json, logging, os, re, subprocess, sys, zipfile
from collections import Counter
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---------- Optional deps ----------
try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None; RGBColor = None; Pt = None; WD_ALIGN_PARAGRAPH = None

try:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos
except Exception:
    FPDF = None; XPos = None; YPos = None

try:
    from ebooklib import epub
except Exception:
    epub = None

# ---------- Logging ----------
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger("fae-v4")

# ---------- Defaults ----------
DEFAULT_SOURCE_DIR = "wiki_pages"
DEFAULT_OUTPUT_DIR = "Federation_Outputs"
DEFAULT_PROJECT_NAME = "Pocket Hunting Dimension"
DEFAULT_EDITION_NAME = "Galactic Blue Edition"
DEFAULT_AUTHORITY = "Saint Council, Year 3960 FS"

# ---------- Colors ----------
COLOR_NAME_TO_RGB = {
    "blue": (0, 94, 184),
    "green": (0, 176, 80),
    "red": (192, 0, 0),
    "purple": (112, 48, 160),
    "gold": (255, 192, 0),
    "gray": (120, 120, 120),
}
COLOR_NAME_TO_HEX = {k: f"#{r:02x}{g:02x}{b:02x}" for k, (r, g, b) in COLOR_NAME_TO_RGB.items()}

# ---------- Taxonomy & realms ----------
CATEGORIES: Dict[str, List[str]] = {
    "Character": [
        "lu ze","nangong jing","qiuyue","alice","lu li","lin ling","yingying",
        "saint","elder","academy","student","cadet","tutor",
        "xuan yuqi","louisa","margaret","saint lin","derrick","tracy",
        "blade demon","dark metal","smoke demon","bakar","goros","santa","harold",
        "elven","spirit","queen","lilith","liria","fey",
        "yaren","yemeng","zhen","luo","tian","general","captain"
    ],
    "Race": ["race","elf","elven","demon","insectoid","spirit","human","blade","hive"],
    "Faction": ["empire","federation","council","academy","sect","union"],
    "Cultivation": ["cultivation","realm","stage","mortal","singularity"],
    "Power": ["god art","divine art","martial","technique","skill"],
    "Artifact": ["item","equipment","armor","artifact","weapon","treasure"],
    "Misc": ["timeline","miscellaneous","information","history","record"]
}
REALM_ORDER = [
    "Mortal","Aperture","Core","Planetary","Star",
    "Cosmic Cloud","Cosmic System","Cosmic Lord",
    "Cosmic Monarch","Emperor","Singularity"
]
FACTION_COLOUR = {
    "human": "blue","federation": "blue",
    "elf": "green","elven": "green",
    "demon": "red",
    "insect": "purple","insectoid": "purple","hive": "purple",
    "spirit": "gold"
}

# ---------- Model ----------
@dataclass
class Entry:
    title: str
    category: str
    story_start_realm: Optional[str]
    peak_era_realm: Optional[str]
    progression_bar: str
    color: str
    content: str
    source_file: str
    url: Optional[str] = None
    synopsis: Optional[str] = None
    attributes: Dict[str, str] = None
    sections: Dict[str, List[str]] = None

# ---------- Utils ----------
def rgb_color(name: str):
    if RGBColor is None: return None
    r,g,b = COLOR_NAME_TO_RGB.get(name, COLOR_NAME_TO_RGB["gray"])
    return RGBColor(r,g,b)

def clean_text(text: str) -> str:
    text = text.replace("\u200b","").replace("\xa0"," ")
    return "\n".join(l.rstrip() for l in text.splitlines())

def compile_word_regex(phrases: List[str]) -> List[re.Pattern]:
    return [re.compile(rf"\b{re.escape(p)}\b", re.IGNORECASE) for p in phrases]

CATEGORY_REGEX = {cat: compile_word_regex(list(set(keys))) for cat, keys in CATEGORIES.items()}
REALM_REGEX = {r: re.compile(rf"\b{re.escape(r)}\b", re.IGNORECASE) for r in REALM_ORDER}

# --- stronger noise/dupe filtering ---
NOISE_PREFIXES = [
    "Take a look at:","You can also be part of the larger Fandom family",
    "Guides","All Help articles","Managing your new community",
    "Getting Started","Rules of this wiki","Original","Other",
    "Luzeoverview","Navigation","External Links"
]
NOISE_PATTERNS = [
    r"^Need help building out this community\??$",
    r"^How to Contribute$",
    r"^Guides?$", r"^Navigation$", r"^External Links?$", r"^Contents?$",
    r"^\[\s*\]$",
    r"^TITLE$", r"^URL$",
    r"^Cultivation\s*$"
]
NOISE_RE = [re.compile(pat, re.IGNORECASE) for pat in NOISE_PATTERNS]

def is_noise(line: str) -> bool:
    s = line.strip()
    if not s: return True
    if any(s.lower().startswith(p.lower()) for p in NOISE_PREFIXES): return True
    if any(rx.match(s) for rx in NOISE_RE): return True
    return False

def dedupe_lines(lines: List[str]) -> List[str]:
    seen, out = set(), []
    for l in lines:
        k = re.sub(r"\s+", " ", l.strip())
        if not k or k in seen: continue
        seen.add(k); out.append(l)
    return out

def split_kv(line: str) -> Optional[Tuple[str,str]]:
    if ":" not in line: return None
    k,v = line.split(":",1)
    k = k.strip().rstrip("[]").replace("Eye  color","Eye color").replace("Hair  color","Hair color")
    v = v.strip().strip("[]")
    if not k or not v: return None
    return k,v

# --- cultivation normalization ---
CULT_STAGE_MAP = {
    r"\b(aperture|apertures?|aperture opening)\b": "Aperture",
    r"\bcore\b": "Core",
    r"\bplanetary\b": "Planetary",
    r"\bstar( state)?\b": "Star",
    r"\bcosmic\s*cloud\b": "Cosmic Cloud",
    r"\bcosmic\s*system\b": "Cosmic System",
    r"\bcosmic\s*lord\b": "Cosmic Lord",
    r"\bcosmic\s*monarch\b": "Cosmic Monarch",
    r"\bemperor\b": "Emperor",
    r"\bsingularity\b": "Singularity",
}
CULT_REGEX = [(re.compile(k, re.IGNORECASE), v) for k,v in CULT_STAGE_MAP.items()]

def normalize_cultivation(text: str) -> Tuple[Optional[str], Optional[str]]:
    found = []
    for rx, canon in CULT_REGEX:
        if rx.search(text):
            found.append((REALM_ORDER.index(canon), canon))
    if not found:
        return None, None
    found = sorted(set(found))
    return found[0][1], found[-1][1]

def realm_progress(text: str) -> Tuple[Optional[str], Optional[str], str]:
    start, peak = normalize_cultivation(text)
    if not start and not peak:
        return None, None, "▱"*12
    if start and not peak:
        peak = start
    i_peak = REALM_ORDER.index(peak)
    filled = int(((i_peak+1)/len(REALM_ORDER))*12)
    return start, peak, "▰"*filled + "▱"*(12-filled)

def faction_color(text: str) -> str:
    t = text.lower()
    for k,v in FACTION_COLOUR.items():
        if re.search(rf"\b{re.escape(k)}\b", t): return v
    return "gray"

def safe_title_from_filename(filename: str) -> str:
    return Path(filename).stem.replace("_"," ").title()

# --- categorization (title/attr aware) ---
CHAR_ATTR_HINTS = {"Age","Gender","Species","Race","Hair color","Eye color",
                   "Height","Weight","Relatives","Spouse","Friends","Occupation","Title"}

def detect_category_scored(text: str, title: str = "") -> str:
    scores = {c:0 for c in CATEGORIES.keys()}
    for cat, pats in CATEGORY_REGEX.items():
        scores[cat] += sum(len(p.findall(text)) for p in pats)

    t = title.lower()
    if any(n in t for n in ["lu ","lu_","luze","nangong","qiuyue","alice","lin ","jing","lian","ying"]):
        scores["Character"] += 5
    if "technique" in t or "items" in t: scores["Power"] += 3
    if "sect" in t or "federation" in t or "academy" in t or "empire" in t: scores["Faction"] += 3
    if any(k in text for k in CHAR_ATTR_HINTS): scores["Character"] += 3
    if "Cultivation\n" in text or re.search(r"^Cultivation\b", text, re.IGNORECASE|re.MULTILINE):
        scores["Cultivation"] += 2

    pref = ["Character","Faction","Race","Cultivation","Power","Artifact","Misc"]
    best = sorted(scores.items(), key=lambda kv: (-kv[1], pref.index(kv[0])))[0][0]
    return best

# --- skip index/home pages ---
INDEX_TITLES = {"home","main page","pocket hunting dimension wiki","luzeoverview"}
def is_indexish(title: str, url: Optional[str]) -> bool:
    t = title.strip().lower()
    if t in INDEX_TITLES: return True
    if url and "/wiki/Main_Page" in url: return True
    return False

# ---------- Parsing ----------
def parse_wiki_text(raw: str) -> Tuple[Optional[str], Optional[str], Dict[str,str], Dict[str,List[str]]]:
    lines = [l.strip() for l in raw.splitlines()]
    lines = [l for l in lines if not is_noise(l)]
    lines = dedupe_lines(lines)

    url, synopsis = None, None
    attributes: Dict[str,str] = {}
    sections: Dict[str,List[str]] = {}

    # URL/Synopsis (first lines)
    for i,l in enumerate(lines[:15]):
        low = l.lower()
        if low.startswith("url:"): url = l.split(":",1)[1].strip()
        elif low.startswith("synopsis"):
            buf = []
            for j in range(i+1, min(i+12, len(lines))):
                if not lines[j]: break
                buf.append(lines[j])
            if buf: synopsis = " ".join(buf)

    ATTR_KEYS_HINTS = [
        "Vital Status","Age","Gender","Species","Race","Hair color","Eye color","Height","Weight",
        "Constitution","Relatives","Friends","Spouse","Spouse(s)","Enemies","Allies",
        "Current Affiliation","Former Affiliation","Military Status","Homeplanet","Star System",
        "Star Field","Cultivation Rank","Cultivation","Young Duke Title","Title","Occupation","Origin"
    ]
    SECTION_NAMES = [
        "Overview","Appearance","Personality","Cultivation","Cultivation Progress",
        "Techniques","Items","Equipment","Relationships","Allies","Enemies",
        "Notable Events","Trivia","History","Abilities","Titles","Affiliations"
    ]

    # Attributes
    i=0
    while i < len(lines):
        l = lines[i]
        kv = split_kv(l)
        if kv and (kv[0] in ATTR_KEYS_HINTS or len(kv[0]) <= 22):
            attributes[kv[0]] = kv[1]; i += 1; continue
        if any(l.startswith(k) for k in ATTR_KEYS_HINTS) and i+1 < len(lines):
            key = next(k for k in ATTR_KEYS_HINTS if l.startswith(k))
            val = lines[i+1].strip(" :")
            if val and len(val) < 400:
                attributes[key] = val; i += 2; continue
        i += 1

    # Sections
    current, buf = None, []
    def flush():
        nonlocal buf, current
        if current and buf:
            paras = []
            for b in buf:
                if b.startswith("- ") or b.startswith("• "):
                    paras.append("• " + b[2:].strip())
                else:
                    paras.append(b)
            sections[current] = dedupe_lines(paras)
        buf = []

    for l in lines:
        if l in SECTION_NAMES:
            flush(); current = l; continue
        if kv := split_kv(l):
            # ignore lone attr pairs leaking into narrative
            if kv[0] in attributes: continue
        if not is_noise(l): buf.append(l)
    flush()

    return url, synopsis, attributes, sections

# ---------- Read entries ----------
def read_entries(source_dir: Path) -> List[Entry]:
    entries: List[Entry] = []
    if not source_dir.exists():
        log.warning("Source directory %s does not exist; no entries.", source_dir)
        return entries

    files = sorted([p for p in source_dir.iterdir() if p.suffix.lower()==".txt"])
    for p in files:
        try:
            raw = p.read_text(encoding="utf-8", errors="ignore")
        except Exception as ex:
            log.error("Failed to read %s: %s", p.name, ex); continue

        content = clean_text(raw)
        url, synopsis, attrs, secs = parse_wiki_text(content)
        title = safe_title_from_filename(p.name)
        if is_indexish(title, url):
            continue

        cat = detect_category_scored(content, title=title)
        start, peak, bar = realm_progress(content)
        col = faction_color(content)
        entries.append(Entry(
            title=title,
            category=cat,
            story_start_realm=start,
            peak_era_realm=peak,
            progression_bar=bar,
            color=col,
            content=content,
            source_file=p.name,
            url=url, synopsis=synopsis, attributes=attrs or {}, sections=secs or {}
        ))

    cat_order = {c:i for i,c in enumerate(CATEGORIES.keys())}
    entries.sort(key=lambda e: (cat_order.get(e.category, 999), e.title))
    return entries

# ---------- Structured exports ----------
def export_structured(entries: List[Entry], outdir: Path, project: str):
    outdir.mkdir(parents=True, exist_ok=True)
    (outdir / f"{project.replace(' ','_')}_Encyclopaedia.json").write_text(
        json.dumps([asdict(e) for e in entries], indent=2, ensure_ascii=False), encoding="utf-8"
    )
    try:
        import yaml
        (outdir / f"{project.replace(' ','_')}_Encyclopaedia.yaml").write_text(
            yaml.dump([asdict(e) for e in entries], allow_unicode=True), encoding="utf-8"
        )
    except Exception:
        log.warning("PyYAML not available; skipping YAML export.")
    with (outdir / f"{project.replace(' ','_')}_index.csv").open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f); w.writerow(["Title","Category","StoryStart","PeakEra","File"])
        for e in entries:
            w.writerow([e.title,e.category,e.story_start_realm,e.peak_era_realm,e.source_file])

# ---------- DOCX ----------
def export_docx(entries: List[Entry], outdir: Path, project: str, edition: str, authority: str):
    if Document is None:
        log.warning("python-docx not available; skipping DOCX export."); return None
    doc = Document()
    h = doc.add_heading(f"{project} Encyclopaedia", 0)
    if WD_ALIGN_PARAGRAPH: h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"{edition}\nCompiled under authority of {authority}")
    try: p.style = "Intense Quote"
    except Exception: pass
    if WD_ALIGN_PARAGRAPH: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(datetime.now().strftime("Generated %Y-%m-%d %H:%M"))
    doc.add_page_break()

    cat_groups: Dict[str, List[Entry]] = {}
    for e in entries: cat_groups.setdefault(e.category, []).append(e)

    for cat in CATEGORIES.keys():
        group = cat_groups.get(cat, [])
        if not group: continue
        doc.add_heading(cat.upper(), level=1)

        for e in group:
            doc.add_heading(e.title, level=2)
            meta=[]
            if e.story_start_realm: meta.append(f"Start: {e.story_start_realm}")
            if e.peak_era_realm: meta.append(f"Peak: {e.peak_era_realm}")
            if meta:
                doc.add_paragraph("  • " + "   |   ".join(meta))
                pbar = doc.add_paragraph()
                run = pbar.add_run(e.progression_bar)
                col = rgb_color(e.color)
                if col: run.font.color.rgb = col
                if Pt: run.font.size = Pt(12)
            if e.url:
                up = doc.add_paragraph(e.url)
                if Pt and up.runs: up.runs[0].font.size = Pt(9)
            if e.synopsis:
                sp = doc.add_paragraph(e.synopsis)
                if Pt and sp.runs: sp.runs[0].font.size = Pt(10)

            if e.attributes:
                tbl = doc.add_table(rows=0, cols=2)
                for k,v in e.attributes.items():
                    row = tbl.add_row().cells
                    row[0].text = k; row[1].text = v
                doc.add_paragraph("")

            for sname, paras in e.sections.items():
                doc.add_heading(sname, level=3)
                for para in paras:
                    if para.startswith("• "):
                        p = doc.add_paragraph(para[2:])
                        try: p.style = doc.styles["List Bullet"]
                        except Exception: pass
                    else:
                        doc.add_paragraph(para)
            doc.add_paragraph("────────────────────────────")
        doc.add_page_break()

    outpath = outdir / f"{project.replace(' ','_')}.docx"
    doc.save(outpath); return outpath

# ---------- PDF: font + safe paragraph writer ----------
from fpdf.errors import FPDFException

def find_unicode_font() -> Optional[str]:
    env_font = os.environ.get("FAE_UNICODE_FONT")
    if env_font and Path(env_font).exists(): return env_font
    home = Path.home()
    for d in [home/"assets"/"fonts"/"dejavu-sans", home/"assets"/"fonts"/"Dejavu", home/"assets"/"fonts"]:
        if d.exists():
            for name in ["DejaVuSans.ttf","DejaVuSansCondensed.ttf","DejaVuSansMono.ttf","DejaVuSerif.ttf"]:
                p = d / name
                if p.exists(): return str(p)
            for p in d.rglob("*.ttf"):
                n = p.name.lower()
                if ("dejavu" in n or "dejavu-sans" in n or "dejavusans" in n) and ("sans" in n or "serif" in n):
                    return str(p)
    for cand in [r"C:\Windows\Fonts\segoeui.ttf", r"C:\Windows\Fonts\arialuni.ttf", r"C:\Windows\Fonts\seguisym.ttf"]:
        if Path(cand).exists(): return cand
    return None

# --- bullet-proof text output for PDF ---
def _normalize_text_for_pdf(text: str, max_token: int = 60) -> str:
    if not text: return text
    text = text.replace("\t"," ").replace("\u00A0"," ")
    parts = re.split(r"(\s+)", text)
    out=[]
    for p in parts:
        if not p or p.isspace(): out.append(p)
        elif len(p) > max_token:
            chunks = [p[i:i+max_token-10] for i in range(0, len(p), max_token-10)]
            out.append(" ".join(chunks))
        else:
            out.append(p)
    return "".join(out)

def _manual_wrap_and_draw(pdf, text: str, line_h: float):
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin
    words = text.split(" ")
    line = ""
    for w in words:
        probe = (line + " " + w).strip() if line else w
        if pdf.get_string_width(probe) <= usable_w:
            line = probe
        else:
            pdf.set_x(pdf.l_margin)
            pdf.cell(usable_w, line_h, line, ln=1)
            line = w
    if line:
        pdf.set_x(pdf.l_margin)
        pdf.cell(usable_w, line_h, line, ln=1)

def mc(pdf, text: str, line_h: float = 6.0):
    if not text: return
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin
    try:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(usable_w, line_h, text)
    except FPDFException:
        safe = _normalize_text_for_pdf(text)
        try:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable_w, line_h, safe)
        except FPDFException:
            _manual_wrap_and_draw(pdf, safe, line_h)

def export_pdf(entries: List[Entry], outdir: Path, project: str, edition: str):
    if FPDF is None:
        log.warning("fpdf2 not available; skipping PDF export."); return None

    unicode_font_path = find_unicode_font()
    pdf = FPDF(format="A4")
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    FAMILY = "Uni" if unicode_font_path else "Helvetica"
    if unicode_font_path:
        pdf.add_font(FAMILY, "", unicode_font_path)

    def setf(size:int): pdf.set_font(FAMILY, size=size)
    def set_small(): setf(10)
    def set_body(): setf(12)

    def header():
        pdf.set_line_width(0.4); pdf.set_draw_color(200,200,200)
        pdf.line(15, 14, 195, 14); pdf.set_y(16)
        set_small(); pdf.set_text_color(120,120,120)
        pdf.cell(0,6,f"{project} — {edition}", align="R"); pdf.set_text_color(0,0,0)
    def footer():
        pdf.set_y(-12); set_small(); pdf.set_text_color(120,120,120)
        pdf.cell(0,6,f"Page {pdf.page_no()}", align="C"); pdf.set_text_color(0,0,0)
    pdf.header = header; pdf.footer = footer

    setf(18)
    if XPos and YPos:
        pdf.cell(0,12,f"{project} Encyclopaedia", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    else:
        pdf.cell(0,12,f"{project} Encyclopaedia", ln=1, align="C")
    pdf.ln(2)

    cat_groups: Dict[str, List[Entry]] = {}
    for e in entries: cat_groups.setdefault(e.category, []).append(e)

    def set_rgb(name: str):
        r,g,b = COLOR_NAME_TO_RGB.get(name, COLOR_NAME_TO_RGB["gray"])
        pdf.set_text_color(r,g,b)

    for cat in CATEGORIES.keys():
        group = cat_groups.get(cat, [])
        if not group: continue
        pdf.set_text_color(0,94,184); setf(14)
        if XPos and YPos: pdf.cell(0,10,cat.upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else: pdf.cell(0,10,cat.upper(), ln=1)
        pdf.set_text_color(0,0,0)

        for e in group:
            pdf.set_draw_color(200,200,200); pdf.set_line_width(0.4)
            pdf.set_fill_color(240,246,252); setf(13)
            if XPos and YPos: pdf.cell(0,8,e.title, new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True)
            else: pdf.cell(0,8,e.title, ln=1, fill=True)

            meta=[]
            if e.story_start_realm: meta.append(f"Start: {e.story_start_realm}")
            if e.peak_era_realm: meta.append(f"Peak: {e.peak_era_realm}")
            if meta:
                set_small(); pdf.set_text_color(80,80,80)
                if XPos and YPos: pdf.cell(0,6,"   |   ".join(meta), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else: pdf.cell(0,6,"   |   ".join(meta), ln=1)
                pdf.set_text_color(0,0,0)
                bar = e.progression_bar if unicode_font_path else e.progression_bar.replace("▰","#").replace("▱","-")
                set_body(); set_rgb(e.color)
                if XPos and YPos: pdf.cell(0,6,bar, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else: pdf.cell(0,6,bar, ln=1)
                pdf.set_text_color(0,0,0)

            if e.url:
                set_small(); pdf.set_text_color(100,100,100)
                mc(pdf, _normalize_text_for_pdf(e.url), line_h=5)
                pdf.set_text_color(0,0,0)

            if e.synopsis:
                set_body()
                mc(pdf, _normalize_text_for_pdf(e.synopsis), line_h=6)
                pdf.ln(1)

            if e.attributes:
                set_small()
                for k,v in e.attributes.items():
                    mc(pdf, _normalize_text_for_pdf(f"{k}: {v}"), line_h=5)
                pdf.ln(1)

            for sname, paras in (e.sections or {}).items():
                setf(12); pdf.set_text_color(0,94,184)
                if XPos and YPos: pdf.cell(0,7,sname, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else: pdf.cell(0,7,sname, ln=1)
                pdf.set_text_color(0,0,0); set_body()
                for para in paras:
                    txt = "• " + para[2:] if para.startswith("• ") else para
                    mc(pdf, _normalize_text_for_pdf(txt), line_h=6)
                pdf.ln(1)
            pdf.ln(2)

    outpath = outdir / f"{project.replace(' ','_')}.pdf"
    pdf.output(str(outpath)); return outpath

# ---------- EPUB ----------
def export_epub(entries: List[Entry], outdir: Path, project: str, edition: str):
    if epub is None:
        log.warning("ebooklib not available; skipping EPUB export."); return None
    book = epub.EpubBook()
    book.set_identifier("FAE-3960"); book.set_title(f"{project}: {edition}"); book.set_language("en")

    # optional font embed
    embedded_font_item = None; font_href = None
    font_path = find_unicode_font()
    if font_path and Path(font_path).exists():
        embedded_font_item = epub.EpubItem(uid="font_main", file_name="fonts/main.ttf",
                                           media_type="application/x-font-ttf",
                                           content=Path(font_path).read_bytes())
        font_href = "fonts/main.ttf"; book.add_item(embedded_font_item)

    css = """
    @namespace epub "http://www.idpf.org/2007/ops";
    body { font-family: serif; line-height: 1.5; margin: 0.8em; color: #111; }
    h2 { font-family: serif; font-size: 1.6em; margin: 0.2em 0 0.2em 0; }
    h3 { font-family: serif; font-size: 1.25em; color: #0b5db8; margin: 0.8em 0 0.3em 0; }
    .meta { font-style: italic; color: #444; margin-bottom: 0.2em; }
    .bar { font-family: monospace; margin: 0.2em 0 0.8em 0; }
    table { border-collapse: collapse; margin: 0.4em 0 0.8em 0; width: 100%; }
    th, td { border-bottom: 1px solid #ddd; padding: 0.2em 0.4em; vertical-align: top; }
    a { color: #0b5db8; text-decoration: none; } a:hover { text-decoration: underline; }
    """
    if font_href:
        css = f"@font-face{{font-family:'MainSerif';src:url('{font_href}');}}\n" \
              "body, h2, h3, th, td { font-family:'MainSerif', serif; }\n" + css

    cstyle = epub.EpubItem(uid="style_nav", file_name="style/style.css", media_type="text/css", content=css)
    book.add_item(cstyle)

    chapters=[]
    for e in entries:
        color_hex = COLOR_NAME_TO_HEX.get(e.color, COLOR_NAME_TO_HEX["gray"])
        parts=[f"<h2>{html.escape(e.title)}</h2>"]
        meta=[]
        if e.story_start_realm: meta.append(f"<b>Start:</b> {html.escape(e.story_start_realm)}")
        if e.peak_era_realm: meta.append(f"<b>Peak:</b> {html.escape(e.peak_era_realm)}")
        if meta:
            parts.append(f"<p class='meta'>{' &nbsp; | &nbsp; '.join(meta)}</p>")
            parts.append(f"<pre class='bar' style='color:{color_hex}'>{html.escape(e.progression_bar)}</pre>")
        if e.url: parts.append(f"<p class='meta'><a href='{html.escape(e.url)}'>{html.escape(e.url)}</a></p>")
        if e.synopsis: parts.append(f"<p>{html.escape(e.synopsis)}</p>")

        if e.attributes:
            parts.append("<table>")
            for k,v in e.attributes.items():
                parts.append(f"<tr><th style='text-align:left;padding-right:0.8em'>{html.escape(k)}</th>"
                             f"<td>{html.escape(v)}</td></tr>")
            parts.append("</table>")

        for sname, paras in (e.sections or {}).items():
            parts.append(f"<h3>{html.escape(sname)}</h3>")
            for ptxt in paras:
                if ptxt.startswith("• "): parts.append(f"<p>• {html.escape(ptxt[2:])}</p>")
                else: parts.append(f"<p>{html.escape(ptxt)}</p>")

        html_doc="\n".join(parts)
        ch = epub.EpubHtml(title=e.title, file_name=f"{e.title.replace(' ','_')}.xhtml", content=html_doc, lang="en")
        ch.add_item(cstyle)
        if embedded_font_item: ch.add_item(embedded_font_item)
        book.add_item(ch); chapters.append(ch)

    book.toc=tuple(chapters); book.spine=["nav"]+chapters
    book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav())
    outpath = outdir / f"{project.replace(' ','_')}.epub"
    epub.write_epub(str(outpath), book); return outpath

# ---------- EPUBCheck ----------
def validate_epub(epub_path: str, outdir: Path, epubcheck_dir: Optional[str] = None) -> bool:
    report_path = outdir / (Path(epub_path).stem + "_epubcheck.txt")
    root = Path(epubcheck_dir).resolve() if epubcheck_dir else Path(os.environ.get("EPUBCHECK_DIR","."))

    if root and (root / "bin" / "epubcheck.bat").exists():
        proc = subprocess.run([str(root/"bin"/"epubcheck.bat"), epub_path],
                              stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, shell=True)
        report_path.write_text(proc.stdout, encoding="utf-8"); return proc.returncode == 0

    if root and (root/"epubcheck.jar").exists() and (root/"lib").exists():
        cp = f"{root/'epubcheck.jar'};{root/'lib'/'*'}"
        proc = subprocess.run(["java","-cp",cp,"com.adobe.epubcheck.tool.Checker",epub_path],
                              stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
        report_path.write_text(proc.stdout, encoding="utf-8"); return proc.returncode == 0

    try:
        if epub is None: raise RuntimeError("ebooklib not installed")
        epub.read_epub(epub_path)
        report_path.write_text("EPUBCheck not found; basic read-back succeeded.\n", encoding="utf-8")
        return True
    except Exception as ex:
        report_path.write_text(f"EPUB basic read-back failed: {ex}\n", encoding="utf-8"); return False

# ---------- Misc ----------
def write_log(entries: List[Entry], outdir: Path):
    with (outdir/"PHD_parse_log.txt").open("w", encoding="utf-8") as f:
        f.write(f"Parse Log – {datetime.now()}\nEntries: {len(entries)}\n")
        for cat,count in Counter(e.category for e in entries).items(): f.write(f"{cat}: {count}\n")

def write_search_script(outdir: Path, project: str):
    s=("import csv,sys,os\nbase=os.path.dirname(__file__)\n"
       f"index=os.path.join(base,'{project.replace(' ','_')}_index.csv')\n"
       "term=' '.join(sys.argv[1:]).lower()\n"
       "with open(index,encoding='utf-8',newline='') as f:\n"
       " r=csv.reader(f);next(r)\n"
       " fnd=[x for x in r if term in ' '.join(x).lower()]\n"
       "print('\\n'.join(f\"[{x[1]}] {x[0]} – {x[4]}\" for x in fnd) or 'No matches.')\n")
    (outdir/"PHD_search_lore.py").write_text(s, encoding="utf-8")

def write_readme(outdir: Path, project: str, edition: str):
    (outdir/"ReadMe.txt").write_text(
        f"{project} Encyclopaedia – {edition}\n"
        "=======================================================\n"
        "Place wiki .txt files in 'wiki_pages'/\n"
        "Outputs: DOCX, PDF, EPUB, JSON, YAML, Index, Search Tool.\n"
        "Usage:\n"
        "  cd Federation_Outputs\n"
        "  python PHD_search_lore.py \"Lu Ze\"\n", encoding="utf-8"
    )

def write_requirements(outdir: Path):
    (outdir/"requirements.txt").write_text("ebooklib\nlxml\npython-docx\nfpdf2\nPyYAML\n", encoding="utf-8")

def zip_outputs(outdir: Path, zip_name: str="Federation_Outputs.zip"):
    zpath = outdir / zip_name
    with zipfile.ZipFile(zpath,"w",zipfile.ZIP_DEFLATED) as z:
        for p in outdir.rglob("*"):
            if p.is_dir() or p.name==zip_name: continue
            z.write(p, arcname=str(p.relative_to(outdir)))
    return zpath

# ---------- CLI / main ----------
def parse_args(argv: List[str]):
    ap = argparse.ArgumentParser(description="Build encyclopaedia outputs from wiki pages.")
    ap.add_argument("--source", default=DEFAULT_SOURCE_DIR)
    ap.add_argument("--output", default=DEFAULT_OUTPUT_DIR)
    ap.add_argument("--project", default=DEFAULT_PROJECT_NAME)
    ap.add_argument("--edition", default=DEFAULT_EDITION_NAME)
    ap.add_argument("--authority", default=DEFAULT_AUTHORITY)
    ap.add_argument("--epubcheck-dir", default=None, help="Folder containing bin/, lib/, epubcheck.jar")
    return ap.parse_args(argv)

def main(argv: List[str] = None):
    args = parse_args(argv or sys.argv[1:])
    source, outdir = Path(args.source), Path(args.output)
    outdir.mkdir(parents=True, exist_ok=True)

    entries = read_entries(source)
    log.info("Entries parsed: %d", len(entries))

    export_structured(entries, outdir, args.project)

    docx_p = export_docx(entries, outdir, args.project, args.edition, args.authority)
    if docx_p: log.info("DOCX: %s", docx_p)

    pdf_p = None
    try:
        pdf_p = export_pdf(entries, outdir, args.project, args.edition)
        if pdf_p: log.info("PDF:  %s", pdf_p)
    except Exception as ex:
        log.warning("PDF export failed but continuing to EPUB: %s", ex)

    epub_p = export_epub(entries, outdir, args.project, args.edition)
    if epub_p:
        log.info("EPUB: %s", epub_p)
        ok = validate_epub(str(epub_p), outdir, epubcheck_dir=args.epubcheck_dir)
        rep = outdir / (epub_p.stem + "_epubcheck.txt")
        log.info("EPUBCheck: %s → %s", "PASS" if ok else "FAIL", rep)

    write_log(entries, outdir)
    write_search_script(outdir, args.project)
    write_readme(outdir, args.project, args.edition)
    write_requirements(outdir)
    zf = zip_outputs(outdir)
    print(f"\n✅ Build complete → {zf}")

if __name__ == "__main__":
    main()
